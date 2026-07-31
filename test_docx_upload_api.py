import requests
import docx
import io

def create_sample_docx() -> bytes:
    doc = docx.Document()
    doc.add_heading("SHAIK MASRADDH - RESUME", 0)
    
    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph("Accomplished AI & Software Engineer with expertise in Machine Learning, RAG systems, and Web applications.")
    
    doc.add_heading("KEY PROJECTS", level=1)
    doc.add_paragraph("Project 1: SoilSmart System - Smart soil monitoring IoT solution using Python & ML models reducing water usage by 35%.")
    doc.add_paragraph("Project 2: Live Chat & Booking Platform - Developed React Native app for 10,000 active users.")
    
    doc.add_heading("TECHNICAL SKILLS", level=1)
    doc.add_paragraph("Python, TypeScript, React, Next.js, FastAPI, SQL, PyTorch, Scikit-Learn, Docker.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def test_docx_pipeline():
    docx_bytes = create_sample_docx()
    filename = "Shaik_Masraddh_Resume.docx"
    
    upload_url = "http://localhost:8000/api/v1/upload"
    files = {"file": (filename, docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    data = {"assistant_id": "interview_coach"}
    
    print(f"1. Uploading '{filename}' to {upload_url}...")
    res = requests.post(upload_url, files=files, data=data)
    print("Upload Status:", res.status_code)
    upload_json = res.json()
    print("Upload Payload:", upload_json)
    
    chat_url = "http://localhost:8000/api/v1/chat"
    query_payload = {
        "assistant_id": "interview_coach",
        "query": "tell me about projects?"
    }
    print(f"\n2. Querying 'tell me about projects?'...")
    res_chat = requests.post(chat_url, json=query_payload)
    print("Chat Status:", res_chat.status_code)
    chat_json = res_chat.json()
    print("Chat Answer:\n", chat_json.get("answer"))
    print("Max Similarity Score:", chat_json.get("max_similarity_score"))

if __name__ == "__main__":
    test_docx_pipeline()
